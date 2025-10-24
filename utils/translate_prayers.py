import pathlib
from pathlib import Path
from docx import Document
from transliterate_manglish import custom_manglish_transliterate
from transliterate_indic import iso_transliterate


def read_docx(file_path: Path):
    """"""
    if not file_path.exists() or not file_path.is_file():
        print(f"The file {file_path} does not exist or is not a file.")
        return None
    doc = Document(file_path.resolve().__str__())
    return doc


def loop_through_docx_folders(file_path: str | Path):
    """Recursively loops through folders and prints their contents.

    Args:
        file_path (str | Path): The path to the file or directory to loop through.
    """
    path = Path(file_path)
    if not path.exists():
        print(f"The path {file_path} does not exist.")
        return

    if path.is_file() and path.suffix == ".docx":
        file_name = path.name
        if "ml_" not in file_name:
            # print(f"Skipping file not in malayalam: {file_name}")
            return

        content = read_docx(path)
        if content is None:
            return
        mn_doc = Document()
        indic_doc = Document()
        for para in content.paragraphs:
            if "link" in para.text.lower() or "section" in para.text.lower():
                mn_doc.add_paragraph(para.text)
                indic_doc.add_paragraph(para.text)
                continue
            mn_doc.add_paragraph(custom_manglish_transliterate(para.text))
            indic_doc.add_paragraph(iso_transliterate(para.text))

        new_path_name = Path(*path.parts[:-1], path.name.replace("ml_", "mn_"))
        mn_doc.save(new_path_name.resolve().__str__())
        new_path_name = Path(*path.parts[:-1], path.name.replace("ml_", "indic_"))
        indic_doc.save(new_path_name.resolve().__str__())
        print(f"Processed file: {file_name}")

    elif path.is_dir():
        # print(f"Directory: {path.name}")
        for item in path.iterdir():
            loop_through_docx_folders(item)


if __name__ == "__main__":
    path = "DOCS/qurbanaSongs"
    loop_through_docx_folders(path)
